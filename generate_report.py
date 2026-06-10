#!/usr/bin/env python3
"""生成符合老师要求的《思想道德与法治》调查报告 Word 文档"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json, os

REPORT_DIR = os.path.dirname(os.path.abspath(__file__))
LOGO_PATH = "/Users/marlin/Downloads/\u6821\u5fbd/\u6821\u5fbd.png"

doc = Document()

# ── 全局样式 ──
style = doc.styles["Normal"]
font = style.font
font.name = "\u5b8b\u4f53"
font.size = Pt(12)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
style.paragraph_format.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)


def add_p(text, size=Pt(12), bold=False, align=None, ls=1.5, indent=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "\u5b8b\u4f53"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
    r.font.size = size
    r.bold = bold
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = ls
    if indent:
        p.paragraph_format.first_line_indent = indent
    return p


def heading(text):
    return add_p(text, bold=True)


def body(text):
    return add_p(text, indent=Cm(0.74))


def att(text, bold=False):
    return add_p(text, size=Pt(10.5), bold=bold, ls=1.0,
                 indent=None if bold else Cm(0.74))


# ── 读取内容数据 ──
with open(os.path.join(REPORT_DIR, "report_content.json"), "r", encoding="utf-8") as f:
    data = json.load(f)

# ── 封面 ──
# 校徽
logo_p = doc.add_paragraph()
logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_r = logo_p.add_run()
logo_r.add_picture(LOGO_PATH, width=Inches(1.8))
add_p("", ls=1.0)
add_p("\u5e7f\u4e1c\u6d77\u6d0b\u5927\u5b66", size=Pt(26), bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER)
for _ in range(2):
    add_p("", ls=2.0)
add_p("\u300a\u601d\u60f3\u9053\u5fb7\u4e0e\u6cd5\u6cbb\u300b\u5b9e\u8df5\u8c03\u67e5\u62a5\u544a",
      size=Pt(22), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_p("", ls=1.5)
add_p("", ls=1.5)

cover = data["cover"]
for label, value in cover:
    t = f"{label}\uff1a{value}" if label else f"               {value}"
    add_p(t, size=Pt(14), align=WD_ALIGN_PARAGRAPH.CENTER, ls=2.0)

doc.add_page_break()

# ── 标题 ──
add_p(data["title"], size=Pt(16), bold=True,
      align=WD_ALIGN_PARAGRAPH.CENTER)
add_p("", ls=1.0)

# ── 团队分工 ──
add_p("\u56e2\u961f\u5206\u5de5\u7b80\u4ecb\uff1a", size=Pt(10.5), bold=True, ls=1.0)
for m in data["team"]:
    add_p(m, size=Pt(10.5), ls=1.0)
add_p("", ls=1.0)

# ── 正文 ──
for sec in data["body"]:
    heading(sec["h"])
    for para in sec["p"]:
        body(para)
    add_p("", ls=1.0)

doc.add_page_break()

# ── 附件一 ──
add_p("\u9644\u4ef6\u4e00\uff1a\u95ee\u5377\u53ca\u63cf\u8ff0\u6027\u7edf\u8ba1\u7ed3\u679c", size=Pt(10.5), bold=True, ls=1.0)
att(data["survey"]["intro"], bold=True)
add_p("", ls=1.0)
att(data["survey"]["greeting"])

for q in data["survey"]["questions"]:
    add_p("", ls=1.0)
    att(q["q"], bold=True)
    # 创建表格：选项 | 人数 | 百分比
    table = doc.add_table(rows=1 + len(q["options"]), cols=3)
    table.style = "Table Grid"
    # 表头
    for ci, ct in enumerate(["选项", "人数", "百分比"]):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(ct)
        r.font.name = "\u5b8b\u4f53"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
        r.font.size = Pt(10.5)
        r.bold = True
    # 数据行
    for ri, opt in enumerate(q["options"]):
        for ci, val in enumerate([opt["text"], str(opt["count"]), opt["pct"]]):
            cell = table.rows[1 + ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "\u5b8b\u4f53"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "\u5b8b\u4f53")
            r.font.size = Pt(10.5)

doc.add_page_break()

# ── 附件二 ──
add_p("\u9644\u4ef6\u4e8c\uff1a\u5b9e\u8df5\u8fc7\u7a0b\u7167\u7247", size=Pt(10.5), bold=True, ls=1.0)
att("\uff08\u6b64\u5904\u63d2\u51654-5\u5f20\u56e2\u961f\u8ba8\u8bba\u3001\u95ee\u5377\u8c03\u67e5\u3001\u6559\u5e08\u6307\u5bfc\u7b49\u573a\u666f\u7684\u7167\u7247\uff09")

doc.add_page_break()

# ── 附件三 ──
add_p("\u9644\u4ef6\u4e09\uff1a\u5b9e\u8df5\u8bf4\u660e", size=Pt(10.5), bold=True, ls=1.0)
for sec in data["notes"]:
    heading(sec["h"])
    for para in sec["p"]:
        body(para)
    add_p("", ls=1.0)

# ── 保存 ──
output = os.path.join(REPORT_DIR, "\u8868\u6f141252\u73ed+\u7f57\u5e94\u6770+\u8c03\u67e5\u62a5\u544a+\u793e\u4f1a\u4e3b\u4e49\u6838\u5fc3\u4ef7\u503c\u89c2.docx")
doc.save(output)
print(f"Done: {output}")
